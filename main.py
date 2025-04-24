import streamlit as st
import tempfile
import os
from pdf2docx import Converter as PDFToDocxConverter
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import time
import zipfile
import base64
from PyPDF2 import PdfReader

# Set page configuration
st.set_page_config(
    page_title="Conversor de PDF y Word",
    page_icon="📕",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .main {padding: 2rem;}
    .stButton button {width: 100%;}
    .stDownloadButton button {width: 100%;}
    h1, h2, h3 {color: #1E88E5;}
    .footer {margin-top: 50px; text-align: center; color: #888;}
    .warning {color: #ff9800; padding: 10px; background-color: #fff3e0; border-radius: 5px; margin: 10px 0;}
    .success {color: #4CAF50; padding: 10px; background-color: #E8F5E9; border-radius: 5px; margin: 10px 0;}
    .info-box {background-color: #f0f7ff; border-left: 5px solid #1E88E5; padding: 15px; margin: 10px 0; border-radius: 5px;}
    .file-item {background-color: #f5f5f5; padding: 10px; margin: 5px 0; border-radius: 4px; display: flex; justify-content: space-between;}
    .pdf-viewer {width: 100%; height: 800px; border: none;}
    </style>
    """, unsafe_allow_html=True)

# Title and description
st.title("Conversor de PDF y Word")
st.markdown("Sube, convierte y descarga múltiples documentos fácilmente.")

# Function to convert PDF to DOCX
def convert_pdf_to_docx(pdf_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        tmp_pdf.write(pdf_file.getvalue())
        pdf_path = tmp_pdf.name

    docx_path = pdf_path.replace('.pdf', '.docx')

    try:
        # Convert PDF to DOCX
        converter = PDFToDocxConverter(pdf_path)
        converter.convert(docx_path)
        converter.close()

        # Read the converted file
        with open(docx_path, 'rb') as f:
            docx_data = f.read()

        # Clean up temporary files
        os.unlink(pdf_path)
        os.unlink(docx_path)

        return docx_data
    except Exception as e:
        st.error(f"Error al convertir PDF a DOCX: {str(e)}")
        # Clean up temporary files
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if os.path.exists(docx_path) and os.path.isfile(docx_path):
            os.unlink(docx_path)
        return None

# Function to convert DOCX to PDF
def convert_docx_to_pdf(docx_data):
    try:
        # Load the Word document
        doc = Document(BytesIO(docx_data))

        # Extract text from the document
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Create a PDF
        buffer = BytesIO()
        pdf_doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        flowables = []
        for text in full_text:
            if text.strip():  # Only add non-empty paragraphs
                flowables.append(Paragraph(text, styles['Normal']))

        pdf_doc.build(flowables)

        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Error al convertir DOCX a PDF: {str(e)}")
        return None

# Create zip file with multiple files
def create_zip_file(files_dict, file_type):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        for file_name, file_data in files_dict.items():
            zip_file.writestr(file_name, file_data)

    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# Function to display PDF content
def display_pdf(pdf_bytes):
    base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" class="pdf-viewer"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

# Function to display DOCX content
def display_docx(docx_bytes):
    try:
        doc = Document(BytesIO(docx_bytes))
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                st.markdown(f"{para.text}")
                if i < len(doc.paragraphs) - 1:
                    st.markdown("---")
    except Exception as e:
        st.error(f"Error al mostrar contenido DOCX: {str(e)}")

# Function to get PDF text content
def get_pdf_text(pdf_bytes):
    try:
        pdf = PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n\n"
        return text
    except Exception as e:
        st.error(f"Error al extraer texto del PDF: {str(e)}")
        return ""

# Initialize session state for file storage
if 'converted_files' not in st.session_state:
    st.session_state.converted_files = {}  # {filename: {data: binary, type: 'docx'/'pdf', timestamp: time}}

if 'selected_file' not in st.session_state:
    st.session_state.selected_file = None

# Sidebar para subida de archivos y opciones de conversión
with st.sidebar:
    # Reemplazar el st.info por una sección de redes sociales llamativa
    st.markdown("### 👇 ¡Conécta conmigo en 1 click! 👇")
    st.markdown('')

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.markdown("""
            <a href="https://www.linkedin.com/in/juan-gonzalo-mart%C3%ADnez-rubio-604316306?lipi=urn%3Ali%3Apage%3Ad_flagship3_profile_view_base_contact_details%3B0XQEwe7NSfmmZkUHkHH8Iw%3D%3D" target="_blank">
                <img src="https://cdn-icons-png.flaticon.com/512/174/174857.png" width="50">
            </a>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
            <a href="https://github.com/Jotaaaaaaa7" target="_blank">
                <img src="https://cdn-icons-png.flaticon.com/512/25/25231.png" width="50">
            </a>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
            <a href="https://www.instagram.com/j00ta._" target="_blank">
                <img src="https://cdn-icons-png.flaticon.com/512/174/174855.png" width="50">
            </a>
        """, unsafe_allow_html=True)

    with col4:
        st.markdown("""
            <a href="mailto:juangonzalo210500@gmail.com" target="_blank">
                <img src="https://cdn-icons-png.flaticon.com/512/5968/5968534.png" width="50">
            </a>
        """, unsafe_allow_html=True)


    st.markdown("---")

    # Uploader unificado - acepta tanto PDF como DOCX
    uploaded_files = st.file_uploader(
        "Sube archivos PDF o Word",
        type=['pdf', 'docx'],
        accept_multiple_files=True
    )

    if uploaded_files:
        st.success(f"{len(uploaded_files)} archivo(s) subido(s) correctamente!")

        # Procesar los archivos subidos según su tipo
        for uploaded_file in uploaded_files:
            file_type = 'pdf' if uploaded_file.name.lower().endswith('.pdf') else 'docx'
            st.session_state.converted_files[uploaded_file.name] = {
                'data': uploaded_file.getvalue(),
                'type': file_type,
                'timestamp': time.time(),
                'original': uploaded_file.name
            }

        # Mostrar botones de conversión por lotes
        if any(file_info['type'] == 'pdf' for file_info in st.session_state.converted_files.values()):
            if st.button("Convertir PDFs a Word"):
                with st.spinner("Convirtiendo PDFs a Word..."):
                    conversion_count = 0
                    for file_name, file_info in list(st.session_state.converted_files.items()):
                        if file_info['type'] == 'pdf':
                            pdf_file = BytesIO(file_info['data'])
                            pdf_file.name = file_name
                            docx_data = convert_pdf_to_docx(pdf_file)
                            if docx_data:
                                docx_name = file_name.replace('.pdf', '.docx')
                                st.session_state.converted_files[docx_name] = {
                                    'data': docx_data,
                                    'type': 'docx',
                                    'timestamp': time.time(),
                                    'original': file_name
                                }
                                conversion_count += 1
                    st.success(f"Conversión completada para {conversion_count} archivo(s)")

        if any(file_info['type'] == 'docx' for file_info in st.session_state.converted_files.values()):
            if st.button("Convertir Words a PDF"):
                with st.spinner("Convirtiendo documentos Word a PDF..."):
                    conversion_count = 0
                    for file_name, file_info in list(st.session_state.converted_files.items()):
                        if file_info['type'] == 'docx':
                            pdf_data = convert_docx_to_pdf(file_info['data'])
                            if pdf_data:
                                pdf_name = file_name.replace('.docx', '.pdf')
                                st.session_state.converted_files[pdf_name] = {
                                    'data': pdf_data,
                                    'type': 'pdf',
                                    'timestamp': time.time(),
                                    'original': file_name
                                }
                                conversion_count += 1
                    st.success(f"Conversión a PDF completada para {conversion_count} documento(s) Word")


    st.markdown("### Acerca de")
    st.info(
        """
        Esta aplicación te permite:
        - Convertir múltiples PDF a documentos Word y al revés
        - Ver el contenido de tus archivos PDF
        - Descargar tus archivos convertidos individualmente o en lote

        """
    )

# Main area for batch download and file management
# Main area for batch download and file management
# Main area for batch download and file management
# Main area for batch download and file management
if st.session_state.converted_files:
    st.header("Archivos Convertidos")

    # Filter options
    file_filter = st.radio("Filtrar por tipo:", ["Todos", "Word (.docx)", "PDF (.pdf)"], horizontal=True)

    filtered_files = {}
    if file_filter == "Word (.docx)":
        filtered_files = {k: v for k, v in st.session_state.converted_files.items() if v['type'] == 'docx'}
    elif file_filter == "PDF (.pdf)":
        filtered_files = {k: v for k, v in st.session_state.converted_files.items() if v['type'] == 'pdf'}
    else:
        filtered_files = st.session_state.converted_files

    if not filtered_files:
        st.info(f"No hay archivos {file_filter.lower()} disponibles.")
    else:
        # Display files as expanders
        for file_name, file_info in filtered_files.items():
            # Calculate file size
            file_size_bytes = len(file_info['data'])
            if file_size_bytes < 1024:
                file_size = f"{file_size_bytes} B"
            elif file_size_bytes < 1024 * 1024:
                file_size = f"{file_size_bytes / 1024:.1f} KB"
            else:
                file_size = f"{file_size_bytes / (1024 * 1024):.1f} MB"

            # Format timestamp
            timestamp = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(file_info['timestamp']))

            # Create expander for each file
            with st.expander(f"{file_name} - {file_info['type'].upper()} - {file_size} - {timestamp}"):
                # Para archivos PDF, usar pestañas
                if file_info['type'] == 'pdf':
                    actions_tab, preview_tab = st.tabs(["Acciones", "Vista previa"])

                    # Pestaña de Acciones
                    with actions_tab:
                        col1, col2 = st.columns(2)

                        with col1:
                            if st.button("Convertir a Word", key=f"convert_pdf_{file_name}", use_container_width=True):
                                with st.spinner(f"Convirtiendo {file_name} a Word..."):
                                    pdf_file = BytesIO(file_info['data'])
                                    pdf_file.name = file_name
                                    docx_data = convert_pdf_to_docx(pdf_file)
                                    if docx_data:
                                        docx_name = file_name.replace('.pdf', '.docx')
                                        st.session_state.converted_files[docx_name] = {
                                            'data': docx_data,
                                            'type': 'docx',
                                            'timestamp': time.time(),
                                            'original': file_name
                                        }
                                        st.success(f"Conversión completada: {docx_name}")
                                        st.rerun()

                        with col2:
                            if st.button("Eliminar archivo", key=f"delete_{file_name}", use_container_width=True):
                                del st.session_state.converted_files[file_name]
                                st.success(f"Archivo {file_name} eliminado")
                                st.rerun()

                        download_name = st.text_input("Nombre para descargar:",
                                                      value=file_name,
                                                      key=f"rename_{file_name}")

                        st.download_button(
                            "Descargar PDF",
                            data=file_info['data'],
                            file_name=download_name,
                            mime="application/octet-stream",
                            key=f"download_{file_name}",
                            use_container_width=True
                        )

                    # Pestaña de Vista previa
                    with preview_tab:
                        display_pdf(file_info['data'])


                # Para archivos Word, mostrar solo acciones (sin pestañas)
                else:
                    col1, col2 = st.columns(2)

                    with col1:
                        if st.button("Convertir a PDF", key=f"convert_docx_{file_name}", use_container_width=True):
                            with st.spinner(f"Convirtiendo {file_name} a PDF..."):
                                pdf_data = convert_docx_to_pdf(file_info['data'])
                                if pdf_data:
                                    pdf_name = file_name.replace('.docx', '.pdf')
                                    st.session_state.converted_files[pdf_name] = {
                                        'data': pdf_data,
                                        'type': 'pdf',
                                        'timestamp': time.time(),
                                        'original': file_name
                                    }
                                    st.success(f"Conversión completada: {pdf_name}")
                                    st.rerun()

                    with col2:
                        if st.button("Eliminar archivo", key=f"delete_{file_name}", use_container_width=True):
                            del st.session_state.converted_files[file_name]
                            st.success(f"Archivo {file_name} eliminado")
                            st.rerun()

                    download_name = st.text_input("Nombre para descargar:",
                                                  value=file_name,
                                                  key=f"rename_{file_name}")

                    st.download_button(
                        "Descargar WORD",
                        data=file_info['data'],
                        file_name=download_name,
                        mime="application/octet-stream",
                        key=f"download_{file_name}",
                        use_container_width=True
                    )

        # Add batch download option (single button)
        st.markdown("---")
        st.subheader("Descargar en lote")

        # Determine which files to include based on current filter
        if file_filter == "Word (.docx)":
            download_files = {k: v['data'] for k, v in filtered_files.items() if v['type'] == 'docx'}
            file_desc = "Word"
            filename = "documentos_word.zip"
        elif file_filter == "PDF (.pdf)":
            download_files = {k: v['data'] for k, v in filtered_files.items() if v['type'] == 'pdf'}
            file_desc = "PDF"
            filename = "documentos_pdf.zip"
        else:
            # When "Todos" is selected, include all files
            download_files = {k: v['data'] for k, v in filtered_files.items()}
            file_desc = "seleccionados"
            filename = "documentos.zip"

        if download_files:
            st.download_button(
                f"Descargar todos los archivos {file_desc}",
                data=create_zip_file(download_files, file_desc.lower()),
                file_name=filename,
                mime="application/zip",
                use_container_width=True
            )
        else:
            st.info(f"No hay archivos {file_desc} para descargar.")
else:
    st.info("Por favor, sube uno o más documentos desde el panel lateral para comenzar.")



# Footer
st.markdown("""
<div class="footer">
    <p>🎈 Desarrollado con Streamlit 🎈</p>
    <p>© 2025 Conversor PDF y Word 📕📘</p>
</div>
""", unsafe_allow_html=True)